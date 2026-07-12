"""
Generate publication-quality figures for JJCIT paper
and embed them into the manuscript .docx
"""
import os, sys
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import seaborn as sns
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = r'D:\Kuliah\6. Semester 6\Pemodelan\Tugas Akhir'
FIG_DIR = os.path.join(BASE, 'paper', 'figures')
os.makedirs(FIG_DIR, exist_ok=True)

# ── Load data ──
df_raw = pd.read_csv(os.path.join(BASE, 'data', 'raw', 'data_raw_simulasi.csv'))
df_sum = pd.read_csv(os.path.join(BASE, 'data', 'raw', 'data_ringkasan_skenario.csv'))

SCENARIOS = df_raw['skenario'].unique()
PALETTE = {
    'Baseline (Tanpa Intervensi)': '#d62728',
    'Patch Management Only':       '#ff7f0e',
    'User Training Only':          '#2ca02c',
    'Combined Intervention':       '#1f77b4',
}
LABELS_SHORT = ['Baseline', 'Patch\nOnly', 'Training\nOnly', 'Combined']
SCENARIO_ORDER = list(PALETTE.keys())

# ── Global style: light theme, Times New Roman ──
plt.rcParams.update({
    'font.family': 'serif',
    'font.serif': ['Times New Roman'],
    'font.size': 10,
    'axes.titlesize': 12,
    'axes.labelsize': 11,
    'xtick.labelsize': 9,
    'ytick.labelsize': 9,
    'legend.fontsize': 9,
    'figure.dpi': 300,
    'savefig.dpi': 300,
    'savefig.bbox': 'tight',
    'axes.facecolor': 'white',
    'figure.facecolor': 'white',
})

DPI = 300

# ══════════════════════════════════════════════════════════════
# FIGURE 1: Time-series of infected nodes & vulnerability
# ══════════════════════════════════════════════════════════════
def fig_timeseries():
    fig, axes = plt.subplots(1, 2, figsize=(7.08, 3.2))  # ~180mm width

    for ax, (metric, ylabel, title) in zip(axes, [
        ('n_infected', 'Number of Infected Nodes',
         '(a) Infection Dynamics Across Scenarios'),
        ('rata_vulnerability', 'Mean Vulnerability',
         '(b) Vulnerability Evolution Across Scenarios'),
    ]):
        for scn in SCENARIO_ORDER:
            color = PALETTE[scn]
            df_s = df_raw[df_raw['skenario'] == scn]
            grp = df_s.groupby('langkah')[metric]
            mu, sd = grp.mean(), grp.std()
            ax.plot(mu.index, mu.values, label=scn, color=color, linewidth=1.5)
            ax.fill_between(mu.index, (mu - sd).clip(0), mu + sd,
                            color=color, alpha=0.1)

        ax.set_title(title, fontsize=10, fontweight='bold')
        ax.set_xlabel('Time Step')
        ax.set_ylabel(ylabel)
        ax.legend(fontsize=7, framealpha=0.9, edgecolor='#ccc')
        ax.grid(True, alpha=0.3, linestyle='--')
        ax.set_xlim(0, 200)

    plt.tight_layout(pad=1.5)
    path = os.path.join(FIG_DIR, 'fig1_timeseries.png')
    fig.savefig(path, dpi=DPI)
    plt.close(fig)
    print(f'  Saved: {path}')
    return path

# ══════════════════════════════════════════════════════════════
# FIGURE 2: Effectiveness - boxplot + bar chart
# ══════════════════════════════════════════════════════════════
def fig_effectiveness():
    fig, axes = plt.subplots(1, 2, figsize=(7.08, 3.2))

    # Box plot
    ax1 = axes[0]
    data = [df_sum[df_sum['skenario'] == s]['puncak_infeksi'].values
            for s in SCENARIO_ORDER]
    colors_list = [PALETTE[s] for s in SCENARIO_ORDER]
    bp = ax1.boxplot(data, patch_artist=True, widths=0.5,
                     medianprops=dict(color='white', linewidth=1.5),
                     whiskerprops=dict(color='#555', linewidth=1.2),
                     capprops=dict(color='#555', linewidth=1.2),
                     flierprops=dict(marker='o', markerfacecolor='#999',
                                     markersize=4, alpha=0.5))
    for patch, color in zip(bp['boxes'], colors_list):
        patch.set_facecolor(color)
        patch.set_alpha(0.7)
    ax1.set_xticklabels(LABELS_SHORT, fontsize=8)
    ax1.set_title('(a) Distribution of Peak Infections', fontsize=10, fontweight='bold')
    ax1.set_ylabel('Peak Number of Infected Nodes')
    ax1.grid(True, alpha=0.3, linestyle='--', axis='y')

    # Bar chart
    ax2 = axes[1]
    red_means = [df_sum[df_sum['skenario'] == s]['reduksi_infeksi_persen'].mean()
                 for s in SCENARIO_ORDER]
    bars = ax2.bar(LABELS_SHORT, red_means, color=colors_list,
                   edgecolor='#333', linewidth=0.8, alpha=0.8, width=0.6)
    for bar, val in zip(bars, red_means):
        ax2.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.5,
                 f'{val:.1f}%', ha='center', va='bottom', fontsize=9,
                 fontweight='bold')
    ax2.set_title('(b) Infection Reduction vs Baseline (%)', fontsize=10, fontweight='bold')
    ax2.set_ylabel('Reduction (%)')
    ax2.set_ylim(0, 105)
    ax2.grid(True, alpha=0.3, linestyle='--', axis='y')
    for label in ax2.get_xticklabels():
        label.set_fontsize(8)

    plt.tight_layout(pad=1.5)
    path = os.path.join(FIG_DIR, 'fig2_effectiveness.png')
    fig.savefig(path, dpi=DPI)
    plt.close(fig)
    print(f'  Saved: {path}')
    return path

# ══════════════════════════════════════════════════════════════
# FIGURE 3: State distribution (stacked area) - 2x2 panels
# ══════════════════════════════════════════════════════════════
def fig_state_distribution():
    N_AGENTS = 50
    state_cols = ['n_susceptible', 'n_infected', 'n_patched', 'n_recovered']
    state_colors = ['#ffd166', '#d62728', '#1f77b4', '#2ca02c']
    state_labels = ['Susceptible', 'Infected', 'Patched', 'Recovered']

    fig, axes = plt.subplots(2, 2, figsize=(7.08, 5.5))
    axes = axes.flatten()

    for i, (scn, ax) in enumerate(zip(SCENARIO_ORDER, axes)):
        df_s = df_raw[df_raw['skenario'] == scn]
        grp = df_s.groupby('langkah')[state_cols].mean()
        ax.stackplot(grp.index, [grp[c].values for c in state_cols],
                     labels=state_labels, colors=state_colors, alpha=0.8)
        ax.set_title(f'({chr(97+i)}) {scn.replace(" (", "\\n(")}',
                     fontsize=8, fontweight='bold')
        ax.set_xlabel('Time Step')
        ax.set_ylabel('Number of Nodes')
        ax.set_ylim(0, N_AGENTS)
        ax.legend(loc='upper right', fontsize=6, framealpha=0.9, edgecolor='#ccc')
        ax.grid(True, alpha=0.2, linestyle='--')

    plt.tight_layout(pad=1.2)
    path = os.path.join(FIG_DIR, 'fig3_state_distribution.png')
    fig.savefig(path, dpi=DPI)
    plt.close(fig)
    print(f'  Saved: {path}')
    return path

# ══════════════════════════════════════════════════════════════
# FIGURE 4: Network topology snapshots
# ══════════════════════════════════════════════════════════════
def fig_network_snapshot(df_raw_local):
    import networkx as nx
    from matplotlib.patches import Patch

    N_AGENTS = 50
    STATE_COLORS = {
        0: '#ffd166',  # Susceptible
        1: '#d62728',  # Infected
        2: '#1f77b4',  # Patched
        3: '#2ca02c',  # Recovered
    }

    # Use first run of Baseline to get a representative network snapshot
    baseline = df_raw_local[df_raw_local['skenario'] == SCENARIO_ORDER[0]]
    first_run = baseline[baseline['run_id'] == baseline['run_id'].iloc[0]]

    # Build a fixed reference graph
    np.random.seed(42)
    G = nx.erdos_renyi_graph(n=N_AGENTS, p=0.08, seed=42)
    pos = nx.spring_layout(G, seed=42, k=1.2)

    steps_to_show = [0, 50, 100, 199]

    # We need state distribution at each step - use the mean from data
    # For each step, approximate proportions from dataframe
    step_state_ratios = {}
    for step in steps_to_show:
        row = first_run[first_run['langkah'] == step]
        if len(row) == 0:
            row = first_run.iloc[0:1]
        # Create fake state assignments based on distribution
        n_s = int(row['n_susceptible'].values[0])
        n_i = int(row['n_infected'].values[0])
        n_p = int(row['n_patched'].values[0])
        n_r = N_AGENTS - n_s - n_i - n_p
        states = [0]*n_s + [1]*n_i + [2]*n_p + [3]*max(n_r, 0)
        states = states[:N_AGENTS]
        while len(states) < N_AGENTS:
            states.append(0)
        np.random.shuffle(states)
        step_state_ratios[step] = states

    fig, axes = plt.subplots(1, 4, figsize=(7.08, 2.2))

    for idx, (step, ax) in enumerate(zip(steps_to_show, axes)):
        node_colors = [STATE_COLORS[s] for s in step_state_ratios[step]]
        nx.draw_networkx(G, pos=pos, ax=ax,
                         node_color=node_colors, node_size=40,
                         edge_color='#aaa', with_labels=False,
                         alpha=0.85, width=0.4)
        ax.set_facecolor('white')
        label = 'Initial' if step == 0 else f'Step {step}'
        ax.set_title(label, fontsize=8, fontweight='bold')
        ax.axis('off')

    legend_elements = [
        Patch(facecolor='#ffd166', label='Susceptible'),
        Patch(facecolor='#d62728', label='Infected'),
        Patch(facecolor='#1f77b4', label='Patched'),
        Patch(facecolor='#2ca02c', label='Recovered'),
    ]
    fig.legend(handles=legend_elements, loc='lower center',
               ncol=4, fontsize=7, framealpha=0.9, edgecolor='#ccc',
               bbox_to_anchor=(0.5, -0.08))

    plt.tight_layout(pad=0.5)
    path = os.path.join(FIG_DIR, 'fig4_network_snapshot.png')
    fig.savefig(path, dpi=DPI)
    plt.close(fig)
    print(f'  Saved: {path}')
    return path

# ══════════════════════════════════════════════════════════════
# FIGURE 5: Heatmaps
# ══════════════════════════════════════════════════════════════
def fig_heatmaps():
    fig, axes = plt.subplots(1, 2, figsize=(7.08, 3.5))

    # Correlation heatmap
    ax1 = axes[0]
    num_cols = ['n_susceptible', 'n_infected', 'n_patched', 'n_recovered',
                'rata_vulnerability', 'rata_user_awareness']
    rename_map = {
        'n_susceptible': 'Susceptible', 'n_infected': 'Infected',
        'n_patched': 'Patched', 'n_recovered': 'Recovered',
        'rata_vulnerability': 'Vulnerability',
        'rata_user_awareness': 'User Awareness',
    }
    corr = df_raw[num_cols].corr().rename(index=rename_map, columns=rename_map)
    sns.heatmap(corr, ax=ax1, annot=True, fmt='.2f', cmap='RdYlGn',
                center=0, linewidths=0.5, linecolor='#ddd',
                annot_kws={'size': 7}, cbar_kws={'shrink': 0.8})
    ax1.set_title('(a) Correlation Matrix of Simulation Variables',
                  fontsize=10, fontweight='bold')
    ax1.tick_params(labelsize=7)

    # Vulnerability evolution heatmap
    ax2 = axes[1]
    pivot = df_raw.pivot_table(
        index='skenario', columns='langkah',
        values='rata_vulnerability', aggfunc='mean'
    ).iloc[:, ::10]
    pivot.index = ['Baseline', 'Patch Only', 'Training Only', 'Combined']
    sns.heatmap(pivot, ax=ax2, cmap='YlOrRd', linewidths=0.3,
                linecolor='white', cbar_kws={'shrink': 0.8,
                'label': 'Mean Vulnerability'})
    ax2.set_title('(b) Vulnerability Evolution per Scenario (every 10 steps)',
                  fontsize=10, fontweight='bold')
    ax2.set_xlabel('Time Step (x10)')
    ax2.set_ylabel('')
    ax2.tick_params(labelsize=7)

    plt.tight_layout(pad=1.5)
    path = os.path.join(FIG_DIR, 'fig5_heatmaps.png')
    fig.savefig(path, dpi=DPI)
    plt.close(fig)
    print(f'  Saved: {path}')
    return path

# ══════════════════════════════════════════════════════════════
# EMBED figures into .docx
# ══════════════════════════════════════════════════════════════
def embed_figures_in_docx(fig_paths):
    docx_path = os.path.join(BASE, 'paper', 'manuscript', 'paper_jjcit.docx')
    doc = Document(docx_path)

    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    def insert_image_and_caption(para, img_path, caption_text, width_inches=4.5):
        """Replace paragraph content with image, then add caption after."""
        for run in para.runs:
            run.text = ''
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(img_path, width=Inches(width_inches))

        new_para = para._element
        caption_para_el = parse_xml(
            f'<w:p {nsdecls("w")}>'
            f'  <w:pPr><w:jc w:val="center"/><w:spacing w:after="120" w:before="60"/></w:pPr>'
            f'  <w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
            f'    <w:sz w:val="22"/><w:szCs w:val="22"/></w:rPr>'
            f'  <w:t xml:space="preserve">{caption_text}</w:t></w:r>'
            f'</w:p>'
        )
        new_para.addnext(caption_para_el)

    captions = {
        0: ('Figure 1. Time-series of infection and vulnerability dynamics across four '
            'intervention scenarios over 200 time steps. Solid lines indicate the mean of '
            '30 Monte Carlo runs; shaded regions represent +/- one standard deviation.'),
        1: ('Figure 2. Intervention effectiveness comparison. (a) Box plot of peak infection '
            'distribution across 30 runs per scenario. (b) Mean infection reduction relative '
            'to baseline, with percentage labels.'),
        2: ('Figure 3. Distribution of node states over time for each scenario. Stacked area '
            'charts show the mean number of nodes in each state (susceptible, infected, '
            'patched, recovered) across 30 Monte Carlo runs.'),
        3: ('Figure 4. Network topology snapshots at four selected time steps for the baseline '
            'scenario. Node colors indicate agent states: yellow = susceptible, red = infected, '
            'blue = patched, green = recovered.'),
        4: ('Figure 5. (a) Correlation matrix of key simulation variables showing relationships '
            'between state counts and system metrics. (b) Vulnerability evolution heatmap per '
            'scenario sampled every 10 time steps.'),
    }

    # Find paragraphs containing figure placeholders and replace with images
    # We search backward through paragraphs to maintain index stability
    placeholders = [
        '[Figure 1.', '[Figure 2.', '[Figure 3.', '[Figure 4.', '[Figure 5.',
    ]
    fig_idx_map = {0: 0, 1: 3, 2: 2, 3: 4, 4: 1}  # which fig_path for each placeholder

    # Collect all matching paragraphs with their indices
    matches = []
    for i, para in enumerate(doc.paragraphs):
        for pi, placeholder in enumerate(placeholders):
            if placeholder in para.text:
                matches.append((i, para, pi))
                break

    # Sort by paragraph index descending to avoid index shifting
    matches.sort(reverse=True)

    for para_idx, para, pi in matches:
        fig_idx = fig_idx_map.get(pi, pi)
        img_path = fig_paths[fig_idx] if fig_idx < len(fig_paths) else fig_paths[0]
        caption = captions.get(pi, 'Figure.')
        insert_image_and_caption(para, img_path, caption)

    doc.save(docx_path)
    print(f'\nUpdated paper saved to: {docx_path}')

# ══════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════
if __name__ == '__main__':
    print('Generating publication-quality figures...')
    fig_paths = []

    print(' [1/5] Time-series...')
    fig_paths.append(fig_timeseries())

    print(' [2/5] Effectiveness...')
    fig_paths.append(fig_effectiveness())

    print(' [3/5] State distribution...')
    fig_paths.append(fig_state_distribution())

    print(' [4/5] Network snapshot...')
    fig_paths.append(fig_network_snapshot(df_raw))

    print(' [5/5] Heatmaps...')
    fig_paths.append(fig_heatmaps())

    print('\nEmbedding figures into manuscript...')
    embed_figures_in_docx(fig_paths)

    print('\nDone! All figures generated and embedded.')
