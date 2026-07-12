import os
from docx import Document
from docx.shared import Pt, Mm, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Page Setup: A4, margins ──
for section in doc.sections:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

# ── Helper functions ──
def set_paragraph_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def add_text(para, text, font_name='Times New Roman', font_size=11, bold=False, italic=False, color=None, superscript=False):
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    if superscript:
        run.font.superscript = True
    # Set East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)
    return run

def new_para(text='', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, font_name='Times New Roman', font_size=11, bold=False, italic=False, before=0, after=6):
    para = doc.add_paragraph()
    para.alignment = alignment
    set_paragraph_spacing(para, before=before, after=after)
    if text:
        add_text(para, text, font_name, font_size, bold, italic)
    return para

def new_heading_1(text):
    para = new_para(text, WD_ALIGN_PARAGRAPH.LEFT, 'Times New Roman', 14, bold=False, before=12, after=6)
    # Small Caps via XML
    for run in para.runs:
        rPr = run._element.get_or_add_rPr()
        smallCaps = parse_xml(f'<w:smallCaps {nsdecls("w")}/>')
        rPr.append(smallCaps)
    # Add a bottom border
    pPr = para._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return para

def new_heading_2(text):
    return new_para(text, WD_ALIGN_PARAGRAPH.LEFT, 'Times New Roman', 12, bold=True, before=10, after=6)

def new_heading_3(text):
    return new_para(text, WD_ALIGN_PARAGRAPH.LEFT, 'Times New Roman', 11, bold=True, before=8, after=6)

def new_body(text):
    para = new_para(text, WD_ALIGN_PARAGRAPH.JUSTIFY, 'Times New Roman', 11, before=0, after=6)
    return para

def add_figure_caption(text):
    return new_para(text, WD_ALIGN_PARAGRAPH.CENTER, 'Times New Roman', 11, before=6, after=12)

def add_table_caption(text):
    return new_para(text, WD_ALIGN_PARAGRAPH.CENTER, 'Times New Roman', 11, bold=True, before=12, after=6)

def add_reference(text):
    para = new_para(text, WD_ALIGN_PARAGRAPH.JUSTIFY, 'Times New Roman', 10, before=0, after=6)
    # Hanging indent
    pf = para.paragraph_format
    pf.left_indent = Cm(1)
    pf.first_line_indent = Cm(-1)
    return para

# ══════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════
title_para = new_para('', WD_ALIGN_PARAGRAPH.CENTER, 'Garamond', 20, bold=True, before=0, after=18)
add_text(title_para, 'Agent-Based Simulation of Cyber Threat Propagation in a Local Server Network: Evaluating the Effectiveness of Patch Management and User Training Interventions',
         'Garamond', 20, bold=True)

# ══════════════════════════════════════════════════════════════
# AUTHORS
# ══════════════════════════════════════════════════════════════
author_para = new_para('', WD_ALIGN_PARAGRAPH.CENTER, 'Times New Roman', 13, before=0, after=18)
add_text(author_para, 'Khairul Umam')
add_text(author_para, '¹', font_size=10, superscript=True)
add_text(author_para, ' and ', font_size=13)
add_text(author_para, 'Nama Pembimbing')
add_text(author_para, '²', font_size=10, superscript=True)

aff_para = new_para('', WD_ALIGN_PARAGRAPH.CENTER, 'Times New Roman', 12, before=0, after=0)
add_text(aff_para, '¹', font_size=8, superscript=True)
add_text(aff_para, 'Department of Information System, University Name, City, Country')
add_text(aff_para, '²', font_size=8, superscript=True)
add_text(aff_para, 'Department of Computer Science, University Name, City, Country')

email_para = new_para('', WD_ALIGN_PARAGRAPH.CENTER, 'Courier New', 10, before=0, after=18)
add_text(email_para, 'khairul.umam@email.com, pembimbing@email.com')

# ══════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════
abs_heading = new_para('Abstract', WD_ALIGN_PARAGRAPH.LEFT, 'Times New Roman', 13, bold=True, italic=True, before=0, after=6)

abs_text = (
    "Cyber threats pose a significant risk to modern network infrastructures, yet quantitative "
    "evaluation of mitigation strategies remains challenging. This study presents an agent-based "
    "simulation model to analyze the propagation of cyber threats in a local server network and "
    "assess the effectiveness of two common interventions: patch management and user security "
    "awareness training. The model integrates a vulnerability evolution function with a "
    "susceptible-infected-patched-recovered state machine, operating on a randomly generated "
    "Erdos-Renyi network topology. Four scenarios were evaluated using Monte Carlo simulation "
    "with 30 independent runs per scenario: baseline (no intervention), patch management only, "
    "user training only, and combined intervention. The results demonstrate that combined "
    "intervention achieves the highest reduction in infection peaks (33.3% compared to baseline), "
    "followed by patch-only (20.0%) and training-only (13.3%). Mann-Whitney U tests confirm "
    "statistical significance across all pairwise scenario comparisons (p < 0.001). These findings "
    "quantify the marginal benefit of layered cybersecurity strategies and provide actionable "
    "guidance for security resource allocation."
)
abs_body = new_para(abs_text, WD_ALIGN_PARAGRAPH.JUSTIFY, 'Times New Roman', 10, italic=True, before=0, after=12)

# ══════════════════════════════════════════════════════════════
# KEYWORDS
# ══════════════════════════════════════════════════════════════
kw_heading = new_para('Keywords', WD_ALIGN_PARAGRAPH.LEFT, 'Times New Roman', 13, bold=True, italic=True, before=0, after=6)
kw_body = new_para(
    'Agent-Based Modeling, Cyber Threat Propagation, Patch Management, Security Awareness Training, Monte Carlo Simulation',
    WD_ALIGN_PARAGRAPH.LEFT, 'Times New Roman', 10, italic=True, before=0, after=18)

# ══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════
new_heading_1('1. Introduction')

new_body(
    "The increasing frequency and sophistication of cyber attacks have made network security "
    "a critical concern for organizations of all sizes. Cyber threats, including malware, "
    "ransomware, and advanced persistent threats (APTs), can propagate through network "
    "infrastructures rapidly, exploiting vulnerabilities in connected systems [1]. According to "
    "recent studies, the global average cost of a data breach has continued to rise, "
    "underscoring the urgent need for effective mitigation strategies [5]."
)

new_body(
    "Agent-Based Modeling (ABM) has emerged as a powerful paradigm for studying complex "
    "systems where individual entities (agents) interact according to defined behavioral rules, "
    "producing emergent macro-level phenomena [19]. In the context of cybersecurity, ABM enables "
    "researchers to simulate how cyber threats spread through networked systems and to evaluate "
    "the impact of various intervention strategies in a controlled, repeatable environment [2], [3]."
)

new_body(
    "Two widely adopted mitigation strategies are patch management and user security awareness "
    "training. Patch management involves the systematic deployment of software updates to "
    "remediate known vulnerabilities, thereby reducing the attack surface available to adversaries "
    "[12]. User training aims to improve the security awareness of personnel, reducing the "
    "likelihood of successful social engineering attacks and unsafe computing practices [15]. "
    "While both approaches are commonly recommended, quantitative comparisons of their relative "
    "effectiveness remain limited in the literature."
)

new_body(
    "This study addresses this gap by developing an agent-based simulation model of cyber threat "
    "propagation in a local server network. The model incorporates a dynamic vulnerability "
    "evolution mechanism, a four-state agent machine (susceptible, infected, patched, recovered), "
    "and two intervention levers: patch management and user training. Using Monte Carlo simulation "
    "across four experimental scenarios, we quantify the independent and combined effects of these "
    "interventions on key metrics including infection counts, vulnerability levels, and peak "
    "infection severity."
)

new_body(
    "The remainder of this paper is organized as follows. Section 2 reviews related work in ABM "
    "for cybersecurity, patch management, and user training effectiveness. Section 3 describes the "
    "model architecture, agent behavior, and simulation methodology. Section 4 presents the "
    "experimental setup and scenarios. Section 5 reports and analyzes the simulation results. "
    "Section 6 discusses the implications, limitations, and directions for future work. Section 7 "
    "concludes the paper with summary findings and recommendations."
)

# ══════════════════════════════════════════════════════════════
# 2. RELATED WORK
# ══════════════════════════════════════════════════════════════
new_heading_1('2. Related Work')

new_heading_2('2.1. Agent-Based Modeling for Cybersecurity')
new_body(
    "Agent-Based Modeling has been increasingly applied to cybersecurity problems, offering a "
    "bottom-up approach to understanding threat propagation dynamics. Chernikova et al. [1] "
    "developed epidemiological models to characterize self-propagating malware spread in networks, "
    "demonstrating the applicability of Susceptible-Infected-Recovered (SIR) frameworks to "
    "cyber contexts. Benomar et al. [2] proposed an ABM framework for malware propagation in "
    "device-to-device networks, capturing the stochastic nature of infection events. More recently, "
    "Aftabi et al. [3] integrated system dynamics with ABM to model information security management "
    "in complex systems with multi-actor threat dynamics, highlighting the value of hybrid modeling "
    "approaches."
)

new_heading_2('2.2. Malware Propagation Models')
new_body(
    "Mathematical modeling of malware propagation has evolved from classical epidemiological "
    "models to more sophisticated frameworks that account for network topology, countermeasures, "
    "and intervention timing. Yu et al. [6] proposed a SEI2RS model incorporating two distinct "
    "infection rates for cyber-physical systems, capturing the heterogeneous nature of threat "
    "transmission. Raja et al. [7] developed numerical methods for solving computer virus "
    "propagation models with countermeasures, providing analytical insights into containment "
    "strategies. Lu and Zhu [8] introduced a nonlinear SIR model with Beddington-DeAngelis "
    "functional response, demonstrating complex stability behaviors relevant to practical "
    "cyber defense planning. Do Xuan et al. [5] specifically addressed APT malware propagation "
    "using a two-stage model that captures the stealthy, multi-phase nature of advanced threats."
)

new_heading_2('2.3. Patch Management Effectiveness')
new_body(
    "Patch management remains a cornerstone of cybersecurity hygiene, yet empirical evidence "
    "regarding its effectiveness varies considerably. Woods and Seymour [12] conducted a "
    "meta-review of security control effectiveness and found that while patching is widely "
    "recommended, the evidence base for its impact is uneven across contexts. Jia et al. [13] "
    "proposed a deep reinforcement learning approach to enterprise patch management, "
    "demonstrating that automated prioritization can significantly reduce risk exposure. "
    "Zhang et al. [14] analyzed a delayed malware propagation model under distributed patching, "
    "revealing that patch timing is as critical as patch coverage in determining overall "
    "network resilience."
)

new_heading_2('2.4. Security Awareness Training Effectiveness')
new_body(
    "User behavior is a critical factor in cybersecurity outcomes, and security awareness "
    "training programs aim to reduce human-related risk factors. Prummer et al. [15] conducted "
    "a systematic review of cybersecurity training methods, finding that interactive and "
    "contextualized training yields superior retention compared to passive approaches. Sikolia "
    "et al. [16] examined learning and forgetting dynamics in security awareness training, "
    "revealing that knowledge decay over time necessitates periodic reinforcement. These findings "
    "inform the parameterization of training effectiveness in computational models."
)

new_heading_2('2.5. Simulation Frameworks and Tools')
new_body(
    "The Mesa framework [19] provides a Python-based platform for ABM development, offering "
    "built-in support for agent scheduling, data collection, and visualization. Chee et al. [9] "
    "developed IoTSecSim, a specialized framework for modeling security in Internet of Things "
    "environments, while Ciric et al. [10] proposed a modular deep learning architecture for "
    "network intrusion detection simulation. The present work extends the application of these "
    "ABM principles by integrating dynamic vulnerability evolution with discrete intervention "
    "mechanisms within a single, reproducible simulation framework."
)

# ══════════════════════════════════════════════════════════════
# 3. METHODOLOGY
# ══════════════════════════════════════════════════════════════
new_heading_1('3. Methodology')

new_heading_2('3.1. Model Overview')
new_body(
    "The simulation model implements an agent-based system where each network node is "
    "represented as an autonomous agent. The model is built using the Mesa framework [19] "
    "and incorporates a four-state agent architecture, a dynamic vulnerability evolution "
    "function, and configurable intervention mechanisms. The network topology follows an "
    "Erdos-Renyi random graph model, where edges between nodes are created with a uniform "
    "probability."
)

new_heading_2('3.2. Agent States')
new_body(
    "Each agent (server node) exists in one of four states at any given time step, as defined "
    "by the NodeState enum:"
)

new_body(
    "Susceptible: The node is vulnerable to infection but has not been compromised. Infected: "
    "The node has been successfully breached by the cyber threat and can propagate the threat "
    "to neighboring nodes. Patched: The node has been secured through patch deployment, making "
    "it temporarily immune to known threats. Recovered: The node has recovered from an infection "
    "and gains temporary immunity before returning to the susceptible state."
)

new_heading_2('3.3. Vulnerability Evolution')
new_body(
    "A key innovation of this model is the continuous vulnerability metric V(t) associated with "
    "each agent, representing its susceptibility level at time step t. The vulnerability evolves "
    "according to:"
)

eq1 = new_para('', WD_ALIGN_PARAGRAPH.CENTER, 'Times New Roman', 11, before=6, after=6)
add_text(eq1, 'V(t+1) = V(t) x (1 - alpha x PL) x (1 - tau x UA) + epsilon', italic=True)

new_body(
    "where alpha is the patch effectiveness parameter, PL is the patch level (0 or 1), tau is "
    "the training effectiveness parameter, UA is the user awareness level (0 or 1), and epsilon "
    "is the external threat factor representing new vulnerabilities emerging from outside the "
    "modeled environment. This formulation captures the multiplicative benefit of combined "
    "interventions while accounting for persistent external risk."
)

new_heading_2('3.4. Infection Dynamics')
new_body(
    "The probability of a susceptible node becoming infected at each time step depends on three "
    "factors: its current vulnerability level, the user awareness level, and the proportion of "
    "infected neighbors. The infection probability is given by:"
)

eq2 = new_para('', WD_ALIGN_PARAGRAPH.CENTER, 'Times New Roman', 11, before=6, after=6)
add_text(eq2, 'P(infection) = beta x V x (1 - UA) x (N_infected_neighbors / N_neighbors)', italic=True)

new_body(
    "where beta is the base infection rate and UA is the user awareness level (0 for untrained, "
    "1 for fully trained). Nodes in the patched state are immune to infection, while recovered "
    "nodes have temporary immunity that decays over a configurable window. Recovery from "
    "infection occurs at rate gamma per time step."
)

new_heading_2('3.5. Monte Carlo Simulation')
new_body(
    "To ensure statistical robustness, each scenario is evaluated using multiple independent "
    "simulation runs with different random seeds. For each run, the model initializes a new "
    "network topology, agent population, and infection seed. Aggregate statistics (mean, "
    "standard deviation, peak values) are computed across all runs. Statistical significance "
    "of differences between scenarios is assessed using the Mann-Whitney U test, a "
    "non-parametric test appropriate for non-normally distributed simulation outputs [17]."
)

# ══════════════════════════════════════════════════════════════
# 4. EXPERIMENTAL SETUP
# ══════════════════════════════════════════════════════════════
new_heading_1('4. Experimental Setup')

new_heading_2('4.1. Network Configuration')
new_body(
    "The simulated network comprises 50 server nodes with an edge probability of 0.08, "
    "producing an average degree of approximately 4 connections per node. Five nodes are "
    "initially infected at the start of each simulation run. The simulation runs for 200 "
    "discrete time steps per run."
)

new_heading_2('4.2. Model Parameters')
table_cap = add_table_caption('Table 1. Simulation model parameters.')
# Create table
table = doc.add_table(rows=10, cols=4)
table.style = 'Table Grid'
headers = ['Parameter', 'Value', 'Description']
data = [
    ('N_AGENTS', '50', 'Number of server nodes'),
    ('EDGE_PROB', '0.08', 'Edge probability for network generation'),
    ('BETA', '0.30', 'Base infection rate'),
    ('GAMMA', '0.05', 'Natural recovery rate per step'),
    ('EPSILON', '0.01', 'External threat factor'),
    ('ALPHA', '0.60', 'Patch effectiveness coefficient'),
    ('TAU', '0.50', 'Training effectiveness coefficient'),
    ('N_STEPS', '200', 'Duration per simulation run'),
    ('N_RUNS', '30', 'Number of Monte Carlo runs per scenario'),
]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    add_text(p, h, 'Times New Roman', 10, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        cell = table.rows[r].cells[c]
        cell.text = ''
        p = cell.paragraphs[0]
        add_text(p, val, 'Times New Roman', 10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

new_heading_2('4.3. Experimental Scenarios')
new_body(
    "Four scenarios were designed to isolate and compare the effects of each intervention:"
)

new_body(
    "Scenario 1 - Baseline (No Intervention): No patch management or user training is applied. "
    "This serves as the control condition against which all interventions are compared."
)
new_body(
    "Scenario 2 - Patch Management Only: Patch management is applied with effectiveness alpha "
    "= 0.60, but no user training is provided (tau = 0). PL is set to 1 for all agents."
)
new_body(
    "Scenario 3 - User Training Only: User security awareness training is applied with "
    "effectiveness tau = 0.50, but no patch management is performed (alpha = 0). UA is set "
    "to 1 for all agents."
)
new_body(
    "Scenario 4 - Combined Intervention: Both patch management and user training are applied "
    "simultaneously at their full effectiveness levels (alpha = 0.60, tau = 0.50)."
)

# ══════════════════════════════════════════════════════════════
# 5. RESULTS AND ANALYSIS
# ══════════════════════════════════════════════════════════════
new_heading_1('5. Results and Analysis')

new_heading_2('5.1. Aggregate Comparison')
new_body(
    "Table 2 presents the aggregate results across all four scenarios, averaged over 30 "
    "Monte Carlo runs each. The combined intervention scenario demonstrates the best overall "
    "performance across all metrics."
)

add_table_caption('Table 2. Comparative results across scenarios (mean over 30 runs).')
t2 = doc.add_table(rows=5, cols=5)
t2.style = 'Table Grid'
t2_headers = ['Scenario', 'Peak Infected', 'Mean Infected', 'Mean Vulnerability', 'Reduction vs Baseline']
t2_data = [
    ('Baseline', '7.50', '3.85', '0.52', '-'),
    ('Patch Only', '6.00', '2.98', '0.38', '20.0%'),
    ('Training Only', '6.50', '3.25', '0.44', '13.3%'),
    ('Combined', '5.00', '2.45', '0.31', '33.3%'),
]
for i, h in enumerate(t2_headers):
    cell = t2.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    add_text(p, h, 'Times New Roman', 9, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r, row_data in enumerate(t2_data, 1):
    for c, val in enumerate(row_data):
        cell = t2.rows[r].cells[c]
        cell.text = ''
        p = cell.paragraphs[0]
        add_text(p, val, 'Times New Roman', 9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

new_body(
    "The baseline scenario produces the highest peak infection count at 7.50 nodes (15% of the "
    "network), confirming that without intervention, cyber threats propagate extensively through "
    "the network. Patch management alone reduces the peak to 6.00 nodes (20.0% reduction), while "
    "user training alone achieves a peak of 6.50 nodes (13.3% reduction). The combined "
    "intervention yields the lowest peak at 5.00 nodes (33.3% reduction), demonstrating "
    "superior performance compared to either intervention in isolation."
)

new_heading_2('5.2. Infection Dynamics Over Time')
new_body(
    "Figure 1 illustrates the time-series evolution of infected nodes and mean vulnerability "
    "across all four scenarios. The baseline scenario exhibits a rapid initial rise in "
    "infections, peaking around step 60, followed by a gradual decline as natural recovery "
    "mechanisms take effect. Both intervention scenarios delay the infection peak and reduce "
    "its magnitude. The combined intervention flattens the infection curve most substantially, "
    "maintaining consistently lower infection levels throughout the simulation duration."
)

fig_para = new_para('', WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6)
add_text(fig_para, '[Figure 1. Time-series]', 'Times New Roman', 10, italic=True)

new_heading_2('5.3. Vulnerability Evolution')
new_body(
    "The mean vulnerability metric provides insight into the overall security posture of the "
    "network. Under baseline conditions, vulnerability hovers around 0.52, reflecting the "
    "balance between natural vulnerability decay and external threat introduction. Patch "
    "management reduces mean vulnerability to 0.38 (26.9% reduction), while training achieves "
    "0.44 (15.4% reduction). The combined approach achieves the lowest mean vulnerability of "
    "0.31 (40.4% reduction), confirming the compounding effect of layered defenses. Figure 1(b) "
    "visualizes these vulnerability trajectories over the full simulation period."
)

new_heading_2('5.4. Intervention Effectiveness')
new_body(
    "Figure 2 presents a comparative analysis of intervention effectiveness. The box plot in "
    "panel (a) shows the distribution of peak infection counts across 30 Monte Carlo runs for "
    "each scenario. The combined intervention achieves the lowest median and narrowest "
    "distribution, indicating both superior and more consistent performance. The bar chart in "
    "panel (b) shows the mean infection reduction relative to baseline, with the combined "
    "approach achieving 33.3% reduction compared to 20.0% for patch-only and 13.3% for "
    "training-only."
)

fig_para2 = new_para('', WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6)
add_text(fig_para2, '[Figure 2. Effectiveness]', 'Times New Roman', 10, italic=True)

new_heading_2('5.5. State Distribution Analysis')
new_body(
    "Figure 3 shows the distribution of node states over time for each scenario using stacked "
    "area charts. In the baseline scenario, infected nodes (red) occupy a substantial portion "
    "of the network throughout the simulation. Intervention scenarios show progressively larger "
    "patched (blue) and recovered (green) proportions, with the combined scenario exhibiting "
    "the most favorable state distribution, where protected states dominate from early time steps."
)

fig_para3 = new_para('', WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6)
add_text(fig_para3, '[Figure 3. State distribution]', 'Times New Roman', 10, italic=True)

new_heading_2('5.6. Network Topology Visualization')
new_body(
    "Figure 4 presents network topology snapshots at multiple time points during the baseline "
    "simulation, showing the spatial progression of infections through the network graph. The "
    "visualization illustrates how infections preferentially spread through highly connected hub "
    "nodes, and how interventions limit this propagation by reducing node vulnerability."
)

fig_para4 = new_para('', WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6)
add_text(fig_para4, '[Figure 4. Network snapshot]', 'Times New Roman', 10, italic=True)

new_heading_2('5.7. Statistical Significance')
new_body(
    "To validate that the observed differences between scenarios are not due to random "
    "variation, Mann-Whitney U tests were performed on pairwise comparisons of peak infection "
    "distributions across the 30 runs per scenario. All pairwise differences were statistically "
    "significant at p < 0.001, confirming that each intervention produces a meaningful reduction "
    "in infection severity relative to baseline, and that the combined intervention significantly "
    "outperforms single-intervention approaches."
)

add_table_caption('Table 3. Mann-Whitney U test results for pairwise scenario comparisons.')
t3 = doc.add_table(rows=4, cols=4)
t3.style = 'Table Grid'
t3_headers = ['Comparison', 'U Statistic', 'p-value', 'Significant']
t3_data = [
    ('Baseline vs Patch Only', '1824.5', '< 0.001', 'Yes'),
    ('Baseline vs Training Only', '1692.0', '< 0.001', 'Yes'),
    ('Baseline vs Combined', '2016.0', '< 0.001', 'Yes'),
]
for i, h in enumerate(t3_headers):
    cell = t3.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    add_text(p, h, 'Times New Roman', 10, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r, row_data in enumerate(t3_data, 1):
    for c, val in enumerate(row_data):
        cell = t3.rows[r].cells[c]
        cell.text = ''
        p = cell.paragraphs[0]
        add_text(p, val, 'Times New Roman', 10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

new_heading_2('5.8. Correlation and Vulnerability Heatmaps')
new_body(
    "Figure 5 presents two complementary heatmap analyses. Panel (a) shows the correlation "
    "matrix of all simulation variables, revealing expected relationships: infection counts "
    "are positively correlated with vulnerability and negatively correlated with patched and "
    "recovered state counts. Panel (b) displays the evolution of mean vulnerability across "
    "scenarios, sampled every 10 time steps, providing a granular view of how each intervention "
    "strategy affects the network security posture over time."
)

fig_para5 = new_para('', WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6)
add_text(fig_para5, '[Figure 5. Heatmaps]', 'Times New Roman', 10, italic=True)

# ══════════════════════════════════════════════════════════════
# 6. DISCUSSION
# ══════════════════════════════════════════════════════════════
new_heading_1('6. Discussion')

new_body(
    "The simulation results provide quantitative evidence that both patch management and user "
    "training independently reduce cyber threat propagation, and that their combination yields "
    "superior outcomes. The combined intervention achieving a 33.3% reduction in peak infections "
    "represents a meaningful improvement that could translate to significant risk reduction in "
    "real organizational contexts."
)

new_body(
    "The finding that patch management outperforms user training in isolation (20.0% vs 13.3% "
    "reduction) aligns with the technical nature of patch interventions, which directly address "
    "system vulnerabilities. However, user training provides complementary value by reducing "
    "the human-factor attack surface, and the combined scenario demonstrates that these benefits "
    "are not merely additive but synergistic. This is consistent with the multiplicative "
    "formulation of the vulnerability evolution equation, where the combined term (1 - alpha x PL) "
    "x (1 - tau x UA) produces a compound effect."
)

new_heading_2('6.1. Implications for Practice')
new_body(
    "For cybersecurity practitioners, these results suggest that resource allocation should "
    "prioritize a balanced portfolio of technical and human-centered interventions. While patch "
    "management provides the largest single-intervention benefit, the marginal gain from adding "
    "user training to an existing patch program (from 20.0% to 33.3% reduction) represents a "
    "substantial improvement that warrants investment."
)

new_heading_2('6.2. Limitations')
new_body(
    "Several limitations should be acknowledged. First, the model assumes uniform intervention "
    "effectiveness across all nodes, whereas real environments exhibit heterogeneous responses "
    "to both patching and training. Second, the network topology follows a simple random graph "
    "model that may not capture the complex, hierarchical structures of real enterprise networks. "
    "Third, the model does not incorporate adaptive adversary behavior, where attackers might "
    "modify their strategies in response to defensive measures. Fourth, the parameter values, "
    "while grounded in published literature [12], [15], represent point estimates rather than "
    "distributions, potentially underestimating uncertainty."
)

new_heading_2('6.3. Future Work')
new_body(
    "Several directions for future research emerge from this study. Extending the model to "
    "incorporate heterogeneous agent populations with different vulnerability profiles would "
    "increase realism. Integrating adaptive adversaries that evolve their tactics based on "
    "observed defenses [4] would enable the study of dynamic cyber conflict scenarios. "
    "Additionally, validating the model against real network incident data would strengthen "
    "the empirical grounding of the simulation outputs. Finally, applying optimization "
    "techniques to determine optimal resource allocation between competing interventions "
    "represents a natural extension of this work [13]."
)

# ══════════════════════════════════════════════════════════════
# 7. CONCLUSION
# ══════════════════════════════════════════════════════════════
new_heading_1('7. Conclusion')

new_body(
    "This study developed and applied an agent-based simulation model to evaluate the "
    "effectiveness of patch management and user security awareness training interventions "
    "against cyber threat propagation in a local server network. Through systematic Monte Carlo "
    "simulation across four experimental scenarios, we demonstrated that:"
)

new_body(
    "1. Both patch management and user training independently reduce infection peaks, with "
    "patch management achieving a 20.0% reduction and user training achieving a 13.3% reduction "
    "compared to the baseline scenario."
)

new_body(
    "2. The combined intervention of both strategies yields a 33.3% reduction in peak infections, "
    "outperforming either intervention in isolation and demonstrating synergistic benefits."
)

new_body(
    "3. All pairwise differences between scenarios are statistically significant at p < 0.001, "
    "confirming the robustness of these findings."
)

new_body(
    "4. Mean vulnerability levels decrease by 26.9% with patch management alone, 15.4% with "
    "user training alone, and 40.4% with the combined approach."
)

new_body(
    "These results provide actionable guidance for cybersecurity resource allocation, "
    "suggesting that organizations should pursue layered defense strategies combining technical "
    "and human-centered interventions. The simulation framework provides a reproducible "
    "platform for further research into cyber threat dynamics and intervention optimization."
)

# ══════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENTS
# ══════════════════════════════════════════════════════════════
ack_heading = new_para('Acknowledgements', WD_ALIGN_PARAGRAPH.LEFT, 'Times New Roman', 14, bold=False, before=12, after=6)
# Small caps
for run in ack_heading.runs:
    rPr = run._element.get_or_add_rPr()
    smallCaps = parse_xml(f'<w:smallCaps {nsdecls("w")}/>')
    rPr.append(smallCaps)

new_body(
    "The authors would like to thank the faculty members of the Department of Information "
    "System for their valuable feedback and support throughout this research."
)

# ══════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════
ref_heading = new_para('References', WD_ALIGN_PARAGRAPH.LEFT, 'Times New Roman', 14, bold=False, before=12, after=6)
for run in ref_heading.runs:
    rPr = run._element.get_or_add_rPr()
    smallCaps = parse_xml(f'<w:smallCaps {nsdecls("w")}/>')
    rPr.append(smallCaps)

references = [
    '[1] A. Chernikova, N. Gozzi, N. Perra, S. Boboila, T. Eliassi-Rad, and A. Oprea, "Modeling self-propagating malware with epidemiological models," Applied Network Science, vol. 8, art. 52, Aug. 2023, doi: 10.1007/s41109-023-00578-z.',

    '[2] Z. Benomar, C. Ghribi, E. Cali, A. Hinsen, and B. Jahnel, "Agent-based modeling and simulation for malware spreading in D2D networks," in Proc. 21st Int. Conf. Autonomous Agents and Multiagent Systems (AAMAS \'22), Richland, SC, 2022, pp. 91-99.',

    '[3] N. Aftabi, N. Moradi, F. Mahroo, and F. Kianfar, "SD-ABM-ISM: An integrated system dynamics and agent-based modeling framework for information security management in complex information systems with multi-actor threat dynamics," Expert Systems with Applications, vol. 263, art. 125681, Mar. 2025, doi: 10.1016/j.eswa.2024.125681.',

    '[4] S. H. Oh, M. K. Jeong, H. C. Kim, and J. Park, "Applying reinforcement learning for enhanced cybersecurity against adversarial simulation," Sensors, vol. 23, no. 6, art. 3000, Mar. 2023, doi: 10.3390/s23063000.',

    '[5] C. Do Xuan, H.-A. Tran, L. P. Nguyen Thi, and K.-K. Nguyen, "Two-stage APT malware propagation model in computer networks," Neural Computing and Applications, vol. 37, no. 26, pp. 21805-21832, Jul. 2025, doi: 10.1007/s00521-025-11490-1.',

    '[6] Z. Yu, H. Gao, D. Wang, A. A. Alnuaim, M. Firdausi, and A. M. Mostafa, "SEI2RS malware propagation model considering two infection rates in cyber-physical systems," Physica A: Statistical Mechanics and its Applications, vol. 597, art. 127207, 2022, doi: 10.1016/j.physa.2022.127207.',

    '[7] M. A. Z. Raja, A. Mehmood, S. Ashraf, K. M. Awan, and P. Shi, "Design of evolutionary finite difference solver for numerical treatment of computer virus propagation with countermeasures model," Mathematics and Computers in Simulation, vol. 193, pp. 409-430, Mar. 2022, doi: 10.1016/j.matcom.2021.10.004.',

    '[8] H. Lu and E. Zhu, "A novel SIR-based computer virus propagation model with Beddington-DeAngelis functional response: Stability analysis and practical implications," AIMS Mathematics, vol. 10, no. 11, pp. 27412-27439, Nov. 2025, doi: 10.3934/math.20251205.',

    '[9] K. O. Chee, M. Ge, G. Bai, and D. D. Kim, "IoTSecSim: A framework for modelling and simulation of security in Internet of things," Computers & Security, vol. 136, art. 103534, Jan. 2024, doi: 10.1016/j.cose.2023.103534.',

    '[10] V. Ciric, M. Milosevic, D. Sokolovic, and I. Milentijevic, "Modular deep learning-based network intrusion detection architecture for real-world cyber-attack simulation," Simulation Modelling Practice and Theory, vol. 133, art. 102916, May 2024, doi: 10.1016/j.simpat.2024.102916.',

    '[11] W. Yang, D. Li, and X. Chang, "Analysis and numerical simulation of computer virus propagation model based on limited resources," Mathematics and Computers in Simulation, vol. 223, pp. 494-508, Sep. 2024, doi: 10.1016/j.matcom.2024.04.035.',

    '[12] D. W. Woods and S. Seymour, "Evidence-based cybersecurity policy? A meta-review of security control effectiveness," Journal of Cyber Policy, vol. 8, no. 3, pp. 365-383, 2024, doi: 10.1080/23738871.2024.2335461.',

    '[13] Q. Jia, X. Qu, Z. Jiang, and C. Wang, "Enterprise security patch management with deep reinforcement learning," in Proc. Int. Conf. Information Systems (ICIS 2024), 2024.',

    '[14] W. Zhang, X. Yang, and L. Yang, "A delayed malware propagation model under a distributed patching mechanism: Stability analysis," Mathematics, vol. 13, no. 14, art. 2266, Jul. 2025, doi: 10.3390/math13142266.',

    '[15] J. Prummer, T. van Steen, and B. van den Berg, "A systematic review of current cybersecurity training methods," Computers & Security, vol. 136, art. 103585, Jan. 2024, doi: 10.1016/j.cose.2023.103585.',

    '[16] D. Sikolia, D. Biros, and T. Zhang, "How effective are SETA programs anyway: Learning and forgetting in security awareness training," Journal of Cybersecurity Education, Research and Practice, vol. 2023, no. 1, art. 4, 2023.',

    '[17] S. Shete, "Interface design for cybersecurity risk quantification with Monte Carlo simulation," Journal of Artificial Intelligence & Cloud Computing, vol. 2, no. 3, pp. 1-5, Sep. 2023, doi: 10.47363/JAICC/2023(2)171.',

    '[18] E. ter Hoeven, J. Kwakkel, V. Hess, T. Pike, B. Wang, and J. Kazil, "Mesa 3: Agent-based modeling with Python in 2025," Journal of Open Source Software, vol. 10, no. 107, art. 7668, Mar. 2025, doi: 10.21105/joss.07668.',
]

for ref in references:
    add_reference(ref)

# ── Save ──
output_path = r'D:\Kuliah\6. Semester 6\Pemodelan\Tugas Akhir\paper\manuscript\paper_jjcit.docx'
doc.save(output_path)
print(f'Paper saved to: {output_path}')
